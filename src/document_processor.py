from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from typing import List, Dict, Any, Union, IO
import os


def extract_text_and_sections(docx_file: Union[str, IO[bytes]], max_paragraphs: int = 10000) -> Dict[str, Any]:
    """
    Extracts text and headings from a .docx file-like object or file path.
    Returns a dict with full_text, headings, and paragraphs.

    - docx_file: file-like object opened in 'rb' or a filesystem path to .docx
    - max_paragraphs: safety cap to avoid processing extremely large documents
    """
    # Allow both file-like objects and file paths
    doc = Document(docx_file)

    paragraphs: List[str] = []
    headings: List[str] = []

    for i, para in enumerate(doc.paragraphs):
        if i >= max_paragraphs:
            break
        text = (para.text or "").strip()
        if not text:
            continue
        # Heuristic: style contains "Heading"
        try:
            if hasattr(para, "style") and para.style and "Heading" in str(para.style.name):
                headings.append(text)
        except Exception:
            # If style access fails for any reason, just skip heading detection
            pass
        paragraphs.append(text)

    full_text = "\n".join(paragraphs)
    return {
        "paragraphs": paragraphs,
        "headings": headings,
        "full_text": full_text
    }


def find_paragraph_indexes(text_snippet: str, paragraphs: List[str]) -> List[int]:
    """
    Returns indexes of paragraphs that contain the snippet (case-insensitive).
    """
    idxs: List[int] = []
    if not text_snippet:
        return idxs
    snippet = text_snippet.lower()
    for i, p in enumerate(paragraphs):
        if snippet in (p or "").lower():
            idxs.append(i)
    return idxs


def annotate_paragraph(doc: Document, para_index: int, comment_text: str, highlight: str = 'YELLOW') -> None:
    """
    Append a COMMENT tag at the end of the paragraph and highlight it.

    - doc: a python-docx Document object (opened separately)
    - para_index: target paragraph index (0-based)
    - comment_text: text to append as a visible comment marker
    - highlight: 'YELLOW' or 'RED' (default YELLOW). Other values fall back to YELLOW.
    """
    try:
        if para_index < 0 or para_index >= len(doc.paragraphs):
            # Out of range, append to last paragraph safely
            para_index = max(0, len(doc.paragraphs) - 1)

        para = doc.paragraphs[para_index]
        run = para.add_run(f" [COMMENT: {comment_text}]")
        run.bold = True

        color = (highlight or 'YELLOW').strip().upper()
        if color == 'RED':
            run.font.highlight_color = WD_COLOR_INDEX.RED
        else:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    except Exception as e:
        print(f"Annotate error at para {para_index}: {e}")


def save_reviewed_copy(original_path: str, doc: Document, output_dir: str = "data/output") -> str:
    """
    Save a reviewed copy with suffix _reviewed.docx and return its path.
    Overwrites existing file with the same name.
    """
    os.makedirs(output_dir, exist_ok=True)
    base = os.path.basename(original_path)
    name, ext = os.path.splitext(base)
    out_path = os.path.join(output_dir, f"{name}_reviewed{ext}")
    doc.save(out_path)
    return out_path
